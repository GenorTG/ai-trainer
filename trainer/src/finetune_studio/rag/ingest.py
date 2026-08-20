"""Document ingestion — parse files and store chunks.

WHAT THIS FILE DOES
==================
The "ETL" of RAG (Extract, Transform, Load):
  1. Extract: read documents (PDF, DOCX, etc.) using parsers.py
  2. Transform: split into chunks of ~512 tokens
  3. Load: embed each chunk and store in the vector database

KEY CONCEPTS
============
- Chunking: large documents are split into smaller pieces because:
  a) LLMs have context limits (can't fit a whole book)
  b) Smaller chunks have more focused meaning
  c) Retrieval is more precise when chunks are focused
- Overlap: consecutive chunks overlap by 50-100 tokens so we don't
  lose information at chunk boundaries.
- Metadata: each chunk remembers where it came from (filename, page,
  chunk index) so we can cite sources.
"""

"""Document ingestion — chunking and embedding for RAG."""
import hashlib
import os
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class Document:
    id: str = ""
    path: str = ""
    filename: str = ""
    content: str = ""
    chunks: list = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    chunk_count: int = 0

@dataclass
class Chunk:
    id: str = ""
    document_id: str = ""
    text: str = ""
    chunk_index: int = 0
    metadata: dict = field(default_factory=dict)


def extract_text(file_path: str) -> str:
    """Extract text from various file formats."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext in (".txt", ".md", ".csv", ".json", ".jsonl", ".py", ".js", ".ts", ".html", ".css"):
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    elif ext == ".pdf":
        return extract_pdf(path)
    elif ext in (".docx", ".doc"):
        return extract_docx(path)
    else:
        # Try reading as text
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        except Exception:  # noqa: BLE001
            return ""


def extract_pdf(path: Path) -> str:
    """Extract text from PDF."""
    try:
        import subprocess
        result = subprocess.run(
            ["pdftotext", str(path), "-"],
            capture_output=True, text=True, timeout=30,
            check=False,
        )
        if result.returncode == 0:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Fallback: PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(path))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except ImportError:
        return f"[PDF: {path.name} — install PyPDF2 or pdftotext to read]"


def extract_docx(path: Path) -> str:
    """Extract text from DOCX."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    except ImportError:
        return f"[DOCX: {path.name} — install python-docx to read]"


def chunk_text(text: str, chunk_size: int = 512, overlap: int = 50, metadata: dict | None = None) -> list[Chunk]:
    """Split text into overlapping chunks."""
    if not text.strip():
        return []

    words = text.split()
    chunks = []
    start = 0
    idx = 0

    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk_words = words[start:end]
        chunk_text_str = " ".join(chunk_words)

        chunks.append(Chunk(
            text=chunk_text_str,
            chunk_index=idx,
            metadata=metadata or {},
        ))
        idx += 1
        start += chunk_size - overlap

    return chunks


def ingest_file(file_path: str, chunk_size: int = 512, overlap: int = 50) -> Document:
    """Ingest a single file — extract text, chunk it."""
    path = Path(file_path)
    content = extract_text(file_path)

    doc_id = hashlib.md5(f"{path.resolve()}".encode()).hexdigest()[:12]

    chunks = chunk_text(content, chunk_size, overlap, {"source": str(path)})

    for i, chunk in enumerate(chunks):
        chunk.document_id = doc_id
        chunk.id = f"{doc_id}_{i}"

    return Document(
        id=doc_id,
        path=str(path.resolve()),
        filename=path.name,
        content=content,
        chunks=chunks,
        chunk_count=len(chunks),
        metadata={"size": path.stat().st_size if path.exists() else 0},
    )


def ingest_directory(directory: str, chunk_size: int = 512, overlap: int = 50,
                     extensions: list | None = None) -> list[Document]:
    """Ingest all supported files in a directory."""
    if extensions is None:
        extensions = [".txt", ".md", ".pdf", ".docx", ".csv", ".json", ".jsonl",
                      ".py", ".js", ".ts", ".html", ".css"]

    documents = []
    for root, dirs, files in os.walk(directory):
        dirs[:] = [d for d in dirs if not d.startswith(".") and d != "__pycache__"]
        for f in sorted(files):
            if any(f.lower().endswith(ext) for ext in extensions):
                fp = os.path.join(root, f)
                doc = ingest_file(fp, chunk_size, overlap)
                if doc.chunk_count > 0:
                    documents.append(doc)

    return documents
